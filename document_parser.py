import io

import pdfplumber
from docx import Document


def extract_text(contents: bytes, file_type: str) -> str:
    if file_type == "PDF":
        with pdfplumber.open(io.BytesIO(contents)) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
            return "\n".join(pages)
    elif file_type == "DOCX":
        doc = Document(io.BytesIO(contents))
        parts = [para.text for para in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)
    return ""
